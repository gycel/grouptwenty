from docx import Document
from docx.shared import Inches
import json

def save_text_and_images_as_docx(text, filename, image_paths=None):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    if image_paths:
        for image_path in image_paths:
            doc.add_picture(image_path, width=Inches(4.5))
    doc.save(filename)

def fill_report_template(template_path, output_path, category_json, status_json, category_chart_img=None, status_chart_img=None, bar_graph_images=None, year_level_names=None):
    doc = Document(template_path)
    category_data = json.loads(category_json)
    status_data = json.loads(status_json)

    # Build replacement map
    replacements = {
        "[no. of Bullying complaints]": str(get_count(category_data, "Bullying")),
        "[no. of Cheating complaints]": str(get_count(category_data, "Cheating")),
        "[no. of Others complaints]": str(get_count(category_data, "Others")),
        "[no. of Teacher-Related complaints]": str(get_count(category_data, "Teacher-Related")),
        "[no. of Pending complaints]": str(get_count(status_data, "Pending")),
        "[no. of On-going complaints]": str(get_count(status_data, "On-going")),
        "[no. of Resolved complaints]": str(get_count(status_data, "Resolved")),
    }

    # Replace in all paragraphs
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    # Insert pie chart images after their respective headings
    def insert_image_after_heading(heading, image_path):
        for i, p in enumerate(doc.paragraphs):
            if heading in p.text and image_path:
                # Insert a new paragraph after the heading and add the image
                new_paragraph = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else doc.add_paragraph()
                run = new_paragraph.add_run()
                run.add_picture(image_path, width=Inches(4.0))
                break

    insert_image_after_heading("Category Distribution", category_chart_img)
    insert_image_after_heading("Status Distribution", status_chart_img)

    # Insert bar graph images after the year level section heading
    if bar_graph_images and year_level_names and len(bar_graph_images) == len(year_level_names):
        for i, p in enumerate(doc.paragraphs):
            if "Number of Complaints per Year Level" in p.text:
                # Move the cursor to after this paragraph
                for img_path, year_name in zip(bar_graph_images, year_level_names):
                    year_paragraph = doc.add_paragraph()
                    year_run = year_paragraph.add_run(year_name)
                    year_paragraph.alignment = 1  # Center
                    year_run.bold = True

                    image_paragraph = doc.add_paragraph()
                    image_run = image_paragraph.add_run()
                    image_run.add_picture(img_path, width=Inches(4.0))
                    image_paragraph.alignment = 1  # Center
                break

    doc.save(output_path)

def get_count(data, key):
    # For category_data: list of dicts with 'category__name' and 'count'
    # For status_data: list of dicts with 'status' and 'count'
    for item in data:
        if 'category__name' in item and item['category__name'] == key:
            return item['count']
        if 'status' in item and item['status'] == key:
            return item['count']
    return 0 